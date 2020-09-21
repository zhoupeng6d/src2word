'''
Date: 2020-09-21 15:42:35
LastEditors: Dash Zhou
LastEditTime: 2020-09-21 17:26:17
'''
from docx import Document
from docx.shared import Cm
import re
import os
import argparse

margin = 2
#path = './src/'

document = Document()

def readFiles(files):
    for i in files:
        document.add_heading(i, 0)
        p = document.add_paragraph('')
        myfile = open(i, 'r',encoding="utf-8")
        lines = myfile.readlines()
        for line in lines:
            if line.split():
                p.add_run(line)

def get_fileList(dir, fileList):
    newDir = dir
    if os.path.isfile(dir):
        if (os.path.splitext(dir)[1] == '.c') or (os.path.splitext(dir)[1] == '.h') or (os.path.splitext(dir)[1] == '.cc'):
            fileList.append(dir)
    elif os.path.isdir(dir):
        for s in os.listdir(dir):
            newDir=os.path.join(dir,s)
            get_fileList(newDir, fileList)
    return fileList

if __name__ =='__main__' :
    parser = argparse.ArgumentParser(description='eg: --path=./src')
    parser.add_argument('--path', type=str, default = './src/')
    args = parser.parse_args()
    print(args.path)

    sections = document.sections

    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)

    list = get_fileList(args.path, [])
    print(len(list))
    for e in list:
        print(e)

    readFiles(list)

    print('========================================')
    print('File output to:' + args.path + '/../output.docx')
    document.save(args.path+'/../output.docx')
